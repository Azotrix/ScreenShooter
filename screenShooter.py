import sys
import os
import pyautogui
import keyboard
import threading
from docx import Document
from datetime import datetime
import subprocess
import platform
import mss
import mss.tools
import msvcrt


# Folder to save screenshots into
TEMP_FOLDER = "screenshots"

# Global listener instance
listener = None

# Global list to store paths of screenshots
screenshots = []

# Path of the final Word document
output_path = ""

# Create an event to control the loop
exit_event = threading.Event()


def capture_screenshot():
    """
    Saves screenshot of the monitor(s) as a (single) PNG file
    Args:
    Returns:
        void:
    """

    # Load global variables
    global screenshots

    os.makedirs(TEMP_FOLDER, exist_ok=True)  # Ensure the folder exists before saving

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    try:
        filepath = os.path.join(TEMP_FOLDER, f"screenshot_{timestamp}.png")
        pyautogui.screenshot(filepath)
        screenshots.append(filepath)
        print(f"Screenshot captured: {filepath}")
    except FileNotFoundError:
        exit_app()


def capture_monitor(index):
    """
    Saves the screen of the monitor with given index into a PNG image
    Args:
        index (int): The index of the monitor to be captured
    Returns:
        void:
    """

    os.makedirs(TEMP_FOLDER, exist_ok=True)  # Ensure the folder exists before saving

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    if monitor_exists(index): 
        with mss.mss() as sct:
            screenshot = sct.grab(sct.monitors[index])
            try:
                filepath = os.path.join(TEMP_FOLDER, f"screenshot_monitor_{index}_{timestamp}.png")
                mss.tools.to_png(screenshot.rgb, screenshot.size, output=filepath)
                screenshots.append(filepath)
                print(f"Screenshot saved for monitor {index}: {filepath}")
            except FileNotFoundError:
                exit_app()
    else:
        capture_screenshot()


def monitor_exists(index):
    """
    Checks whether a monitor with a given index exists
    Args:
        index (int): The index of the sought monitor
    Returns:
        bool: True if a monitor with the given index is found, False otherwise
    """

    with mss.mss() as sct:
        # Check if the index is within the bounds of the monitors list
        return 0 <= index < len(sct.monitors)


def delete_files():
    """
    Removes saved screenshots from the disk
    Args:
    Returns:
        void:
    """

    for screenshot in screenshots:
        try:
            os.remove(screenshot)
        except OSError:
            # Do nothing
            pass

    print("Screenshots deleted.")

    try:
        os.rmdir(TEMP_FOLDER) # Try and remove the folder
    except OSError:
        print(f"The folder '{TEMP_FOLDER}' is not empty, therefore not deleted.")


def save_to_word():
    """
    Saves the created screenshots into a Word document (.docx), and deletes the saved screenshots
    Args:
    Returns:
        void:
    """

    # Load global variables
    global output_path

    if not screenshots:
        print("No screenshots to save.")
        return
    
    print("Saving Word document...")

    doc = Document()
    doc.add_heading("Captured Screenshots", level=1)

    for screenshot in screenshots:
        doc.add_paragraph(f"Screenshot: {os.path.basename(screenshot)}")
        # Calculate the usable width of the page
        section = doc.sections[0]
        usable_width = section.page_width - section.left_margin - section.right_margin

        # Add the picture with the adjusted width
        doc.add_picture(screenshot, width=usable_width)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = os.path.join(os.getcwd(), f"Screenshots_{timestamp}.docx")
    doc.save(output_path)
    print(f"Word document saved: {output_path}")

    delete_files()


def open_word_document():
    """
    Opens the newly created Word document on Windows or Unix based system
    Args:
    Returns:
        void:
    """
        
    try:
        os.startfile(output_path)  # Windows
    except AttributeError:
        os.system(f"open {output_path}")  # macOS/Linux


def confirm_exit():
    """
    Function for prompting the user before exiting
    Args:
    Returns:
        void:
    """

    # Clear all hotkeys to prevent interference with input()
    keyboard.clear_all_hotkeys()
        
    if len(screenshots):
        # Remove previous keystrokes making their way into user input
        while msvcrt.kbhit():
            msvcrt.getch()

        # Prompt the user for confirmation
        response = input("Would you like to delete the screenshots? (Yes/No/Cancel) [Cancel]: ").strip().lower()

        if response in ['yes', 'y', 'ye', 'yeah']: # '' For Enter key
            delete_files()
            exit_app()
        elif response in ['no', 'n', 'nah', 'nope']:
            # Open the folder containing screenshots
            #print(f"Opening the folder: {TEMP_FOLDER}")
            if platform.system() == "Windows":
                os.startfile(TEMP_FOLDER)
            elif platform.system() == "Darwin":  # macOS
                subprocess.run(["open", TEMP_FOLDER])
            else:  # Linux and other Unix-like systems
                subprocess.run(["xdg-open", TEMP_FOLDER])

            exit_app()
        else:
            # Re-register the hotkeys if user cancels
            register_hotkeys()
    else:
        exit_app()


def exit_app():
    """
    Function for exiting the application
    Args:
    Returns:
        void:
    """

    # Stop all keyboard hooks before exiting
    keyboard.unhook_all()
    # Set the event to break the main loop
    exit_event.set()

    sys.exit()

def register_hotkeys():
    """
    Registers all hotkeys for the application
    Args:
    Returns:
        void:
    """
    keyboard.add_hotkey("alt+0", capture_screenshot)
    keyboard.add_hotkey("alt+1", lambda: capture_monitor(1))
    keyboard.add_hotkey("alt+2", lambda: capture_monitor(2))
    keyboard.add_hotkey("alt+3", lambda: capture_monitor(3))
    keyboard.add_hotkey("alt+f10", lambda: (save_to_word(), open_word_document(), exit_app()))
    keyboard.add_hotkey("alt+f12", confirm_exit)
    print("------------")
    print("Press Alt+F12 to exit the application.")
    print("_______________________________")


def main():
    """
    Main function
    """
        
    '''
    # Used for coloured terminals, but does not work for older systems
    print("Press \033[93mAlt+0\033[0m to take a screenshot, \033[92mAlt+F10\033[0m to save them to a Word document.")
    print("For multi-monitor setups, use the keys \033[93mAlt+1\033[0m (1st monitor) / \033[93mAlt+2\033[0m (2nd monitor) / \033[93mAlt+3\033[0m (3rd monitor) to capture the screen of a single monitor (up to 3).")
    print("------------")
    print("Press \033[91mAlt+F12\033[0m to exit the application.")
    print("_______________________________")
    '''
    print("Press Alt+0 to take a screenshot, Alt+F10 to save them to a Word document.")
    print("For multi-monitor setups, use the keys Alt+1 (1st monitor) / Alt+2 (2nd monitor) / Alt+3 (3rd monitor) to capture the screen of a single monitor (up to 3).")

    # Register key combinations
    register_hotkeys()

    try:
        # Run an event-based loop instead of blocking with `keyboard.wait()`
        while not exit_event.is_set():
            exit_event.wait(0.1)  # Check every 100ms
    except KeyboardInterrupt:
        print("\nExiting application...")

main()